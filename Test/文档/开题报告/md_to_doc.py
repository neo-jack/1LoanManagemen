# -*- coding: utf-8 -*-
"""
将开题报告Markdown文件转换为Word文档
需要安装: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re
import os

def set_chinese_font(run, font_name='宋体', font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_title(doc, text, level=1):
    """添加标题"""
    # 清理markdown标记
    text = re.sub(r'^#+\s*', '', text).strip()
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # 移除加粗标记
    
    if level == 1:
        p = doc.add_heading('', level=0)
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        p = doc.add_heading('', level=1)
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 14)
    elif level == 3:
        p = doc.add_heading('', level=2)
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 12)
    elif level == 4:
        p = doc.add_heading('', level=3)
        run = p.add_run(text)
        set_chinese_font(run, '黑体', 12)
    return p

def add_paragraph(doc, text, indent=True):
    """添加正文段落"""
    # 处理加粗文本
    p = doc.add_paragraph()
    
    # 设置段落格式
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    
    # 解析文本中的加粗部分
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # 加粗文本
            run = p.add_run(part[2:-2])
            run.bold = True
            set_chinese_font(run, '宋体', 12)
        else:
            run = p.add_run(part)
            set_chinese_font(run, '宋体', 12)
    
    return p

def add_list_item(doc, text, level=0, original_line=''):
    """添加列表项"""
    # 处理列表项文本
    text = re.sub(r'^[-*]\s*', '', text).strip()
    
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    
    # 根据层级设置缩进和前缀
    if level == 0:
        # 一级列表：无缩进，使用圆点
        p.paragraph_format.first_line_indent = Cm(0.74)
        prefix = ''
    else:
        # 二级及以下列表：有缩进
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.first_line_indent = Cm(0)
        prefix = '- '
    
    # 解析文本中的加粗部分
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    first_part = True
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            content = (prefix if first_part else '') + part[2:-2]
            run = p.add_run(content)
            run.bold = True
            set_chinese_font(run, '宋体', 12)
        else:
            content = (prefix if first_part else '') + part
            run = p.add_run(content)
            set_chinese_font(run, '宋体', 12)
        if first_part and prefix:
            first_part = False
            prefix = ''
    
    return p

def add_numbered_item(doc, text, number):
    """添加编号列表项"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.left_indent = Cm(0.74)
    
    # 解析文本中的加粗部分
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    first = True
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            set_chinese_font(run, '宋体', 12)
        else:
            run = p.add_run(part)
            set_chinese_font(run, '宋体', 12)
        first = False
    
    return p

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    
    # 设置表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                set_chinese_font(run, '宋体', 10)
    
    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = cell_data
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, '宋体', 10)
    
    return table

def parse_md_table(lines, start_idx):
    """解析Markdown表格"""
    headers = []
    rows = []
    
    # 解析表头
    if start_idx < len(lines):
        header_line = lines[start_idx].strip()
        if '|' in header_line:
            headers = [cell.strip() for cell in header_line.split('|') if cell.strip()]
    
    # 跳过分隔行
    idx = start_idx + 2
    
    # 解析数据行
    while idx < len(lines):
        line = lines[idx].strip()
        if not line or not line.startswith('|'):
            break
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
        idx += 1
    
    return headers, rows, idx

def convert_md_to_doc(md_file, doc_file):
    """将Markdown文件转换为Word文档"""
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    i = 0
    in_code_block = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 跳过空行
        if not stripped:
            i += 1
            continue
        
        # 代码块处理
        if stripped.startswith('```'):
            in_code_block = not in_code_block
            i += 1
            continue
        
        if in_code_block:
            i += 1
            continue
        
        # 标题处理
        if stripped.startswith('# '):
            add_title(doc, stripped, 1)
        elif stripped.startswith('## '):
            add_title(doc, stripped, 2)
        elif stripped.startswith('### '):
            add_title(doc, stripped, 3)
        elif stripped.startswith('#### '):
            add_title(doc, stripped, 4)
        
        # 表格处理
        elif stripped.startswith('|') and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers, rows, next_idx = parse_md_table(lines, i)
            if headers and rows:
                add_table(doc, headers, rows)
                doc.add_paragraph()  # 表格后空行
            i = next_idx
            continue
        
        # 编号列表处理
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s*', '', stripped)
            add_numbered_item(doc, text, 0)
        
        # 无序列表处理
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # 计算缩进层级
            leading_spaces = len(line) - len(line.lstrip())
            level = leading_spaces // 2
            add_list_item(doc, stripped, level, line)
        
        # 普通段落
        else:
            add_paragraph(doc, stripped)
        
        i += 1
    
    # 保存文档
    doc.save(doc_file)
    print(f'转换完成: {doc_file}')

if __name__ == '__main__':
    # 获取当前脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 设置输入输出文件路径
    md_file = os.path.join(script_dir, '高校学生助学贷款管理系统开题报告.md')
    doc_file = os.path.join(script_dir, '高校学生助学贷款管理系统开题报告.docx')
    
    # 执行转换
    convert_md_to_doc(md_file, doc_file)
